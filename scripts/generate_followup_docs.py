#!/usr/bin/env python3
"""
Generate follow-up question documents (EN/JA) in .docx format — v3.
All FU-1~FU-10 items from ClientAnswers4.pdf have been implemented.

Includes:
  - Current implementation status for all contract items (tables)
  - Test suite details (33 tests, 233 checks) (tables)
  - FU-1~FU-10 implementation summary (all resolved)
  - Remaining Phase 2 planning questions
  - Page layout: tables and sections do not split across pages
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import date


# ============================================================
# Layout helpers
# ============================================================

HEADER_BG = "4472C4"  # Blue header matching the reference PDF
DONE_BG = "E2EFDA"    # Light green for [DONE]
PARTIAL_BG = "FFF2CC"  # Light yellow for [PARTIAL]
WHITE_BG = "FFFFFF"


def keep_paragraph_together(paragraph):
    """Set keepNext + keepLines so paragraph stays on one page."""
    pPr = paragraph._p.get_or_add_pPr()
    keep_next = parse_xml(f'<w:keepNext {nsdecls("w")} />')
    keep_lines = parse_xml(f'<w:keepLines {nsdecls("w")} />')
    pPr.append(keep_next)
    pPr.append(keep_lines)


def keep_table_on_page(table):
    """Set keepNext + keepLines on every paragraph in every cell
    so the table does not split across pages."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                pPr = paragraph._p.get_or_add_pPr()
                keep_next = parse_xml(f'<w:keepNext {nsdecls("w")} />')
                keep_lines = parse_xml(f'<w:keepLines {nsdecls("w")} />')
                pPr.append(keep_next)
                pPr.append(keep_lines)


def set_cell_bg(cell, color):
    """Set background color on a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear" />'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=Pt(9), color=None, align=None):
    """Set text content of a cell with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def add_heading_keep(doc, text, level=1):
    """Add heading and keep it together with next content."""
    heading = doc.add_heading(text, level=level)
    keep_paragraph_together(heading)
    return heading


def set_table_borders(table):
    """Set thin borders on all cells (matching reference PDF style)."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")} />')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def make_impl_table(doc, headers, rows):
    """Create a 3-column implementation status table (Item | Description | Status).
    Each row: (item_id, description, status_text, status_type).
    status_type: 'done', 'partial', 'config', 'na'.
    """
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, size=Pt(9), color=RGBColor(0xFF, 0xFF, 0xFF))

    # Data rows
    for r_idx, (item_id, desc, status_text, status_type) in enumerate(rows):
        row = table.rows[1 + r_idx]
        set_cell_text(row.cells[0], item_id, bold=True, size=Pt(9))
        set_cell_text(row.cells[1], desc, size=Pt(9))
        set_cell_text(row.cells[2], status_text, size=Pt(9))
        if status_type == "done":
            set_cell_bg(row.cells[2], DONE_BG)
        elif status_type == "partial":
            set_cell_bg(row.cells[2], PARTIAL_BG)
        elif status_type == "config":
            set_cell_bg(row.cells[2], PARTIAL_BG)

    # Column widths
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(10.0)
        row.cells[2].width = Cm(4.5)

    keep_table_on_page(table)
    return table


def make_test_table(doc, test_rows):
    """Create a 3-column test results table (Test # | Description | Result).
    Each row: (test_num, description, result_str).
    """
    headers = ["Test #", "Description", "Result"]
    table = doc.add_table(rows=1 + len(test_rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, size=Pt(9), color=RGBColor(0xFF, 0xFF, 0xFF))

    for r_idx, (tnum, desc, result) in enumerate(test_rows):
        row = table.rows[1 + r_idx]
        set_cell_text(row.cells[0], str(tnum), size=Pt(9))
        set_cell_text(row.cells[1], desc, size=Pt(9))
        set_cell_text(row.cells[2], result, size=Pt(9))

    for row in table.rows:
        row.cells[0].width = Cm(1.8)
        row.cells[1].width = Cm(11.0)
        row.cells[2].width = Cm(2.2)

    keep_table_on_page(table)
    return table


def make_metrics_table(doc, metric_rows):
    """Create a 2-column metrics table (Metric | Description)."""
    headers = ["Metric", "Description"]
    table = doc.add_table(rows=1 + len(metric_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, size=Pt(9), color=RGBColor(0xFF, 0xFF, 0xFF))

    for r_idx, (name, desc) in enumerate(metric_rows):
        row = table.rows[1 + r_idx]
        set_cell_text(row.cells[0], name, bold=True, size=Pt(9))
        set_cell_text(row.cells[1], desc, size=Pt(9))

    for row in table.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(13.0)

    keep_table_on_page(table)
    return table


def make_file_table(doc, file_rows):
    """Create a 3-column file changes table (File | Changes | Lines Modified)."""
    headers = ["File", "Changes", "Lines Modified"]
    table = doc.add_table(rows=1 + len(file_rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, HEADER_BG)
        set_cell_text(cell, h, bold=True, size=Pt(9), color=RGBColor(0xFF, 0xFF, 0xFF))

    for r_idx, (fname, changes, lines) in enumerate(file_rows):
        row = table.rows[1 + r_idx]
        set_cell_text(row.cells[0], fname, bold=True, size=Pt(9))
        set_cell_text(row.cells[1], changes, size=Pt(9))
        set_cell_text(row.cells[2], lines, size=Pt(9))

    for row in table.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(9.0)
        row.cells[2].width = Cm(3.5)

    keep_table_on_page(table)
    return table


def add_question(doc, num, topic, question, context="", priority="Medium"):
    """Add a follow-up question block with keep-together."""
    p = doc.add_paragraph()
    run = p.add_run(f"FU-{num}: {topic}")
    run.bold = True
    run.font.size = Pt(11)
    keep_paragraph_together(p)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Priority: {priority}")
    run2.italic = True
    run2.font.size = Pt(10)
    keep_paragraph_together(p2)

    if context:
        p3 = doc.add_paragraph()
        run3 = p3.add_run(f"Context: {context}")
        run3.font.size = Pt(10)
        keep_paragraph_together(p3)

    p4 = doc.add_paragraph()
    run4 = p4.add_run(f"Question: {question}")
    run4.font.size = Pt(10)
    keep_paragraph_together(p4)

    p5 = doc.add_paragraph()
    run5 = p5.add_run("Answer: ")
    run5.font.size = Pt(10)
    run5.italic = True

    doc.add_paragraph("")


# ============================================================
# Data definitions (shared by EN and JA)
# ============================================================

IMPL_CORE_ARCH = [
    ("C-01", "ReLU Audit Hook (External Wrapper)",
     "SHARK black box; theta_act via external wrapper (Timer 2)", "done"),
    ("C-02", "Cap_c Capability Vector",
     "device_class, cpu_freq, ram_mb, net_type, is_default_cap (Types.hpp)", "done"),
    ("C-03", "3-Tier SLO Filtering",
     "Tier 1: D < e_on (403), Tier 2: D < factor*E_{i,c} (403), Tier 3: VFT (404)", "done"),
    ("C-04", "n_alt Model Negotiation Mapping (FU-6)",
     "Family-based N_ALT_MAP: alexnet:simc2,hinet; batch preserved from original request", "done"),
    ("C-05", "Differentiated EWMA",
     "alpha_pre=0.2 (server-local), alpha_on=0.4 (client-dependent)", "done"),
    ("C-06", "VFT Admission Control",
     "Dynamic feasibility check with VFT_SAFETY_MARGIN=1.2", "done"),
    ("C-07", "Proactive Pre-Processing Buffer",
     "FileManager: CREATING/READY/IN_USE/USED lifecycle", "done"),
]

IMPL_SCHED_ALGO = [
    ("C-08", "DQN Background Training Pipeline",
     "train_dqn.py: 5->64->32->1 MLP, exports DQN2 binary", "done"),
    ("C-09", "7 Scheduler Modes",
     "FCFS, SJF-Aging, EDF, LST, SA, DQN, CUSTOM/SPLS via SchedulerFactory", "done"),
    ("C-10", "Tier 2 Static SLO Filter",
     "SLO >= SLO_FACTOR_TIER2 * E_{i,c} where E_{i,c} = e_pre + theta_c * e_on", "done"),
    ("C-11", "Config-Driven DQN Normalization (FU-1)",
     "TOTAL_CORES=24 (edge node target), TOTAL_MEMORY_MB (DQNScheduler.hpp)", "done"),
    ("C-12", "SA 4D Perturbation",
     "swap_jobs, adjust_malleable_cores, shift_dwell_time, reverse_subsequence", "done"),
    ("C-13~C-16", "SA Multi-Objective Cost Function",
     "Tardiness^2 + 10*MemPenalty + SlackBonus + CoreOvercommit + NegotiationPenalty(phi)", "done"),
    ("C-17", "SA Event-Triggered Re-Optimization",
     "notifyJobComplete() with SA_MIN_RETRIGGER_MS=50", "done"),
    ("C-18", "SA Full-Queue Visibility",
     "SA_WINDOW_SIZE=100 (optimize up to depth 100)", "done"),
    ("C-19", "DQN 5->64->32->1 Architecture",
     "Pointwise ranking MLP with ReLU (DQNScheduler.hpp)", "done"),
    ("C-20", "DQN2 Binary Format + CRC32",
     "28-byte header (magic+dims+arch_hash+crc32) + raw floats", "done"),
]

IMPL_RESOURCE = [
    ("C-21", "Strict Core Allocation",
     "acquireCoresStrict(): exact match or reject (ResourceManager.cpp)", "done"),
    ("C-22", "Dynamic Watchdog with theta_c",
     "timeout_s = min(theta*e_on*1.2/1000+1, MAX_LEASE_TIMEOUT_S)", "done"),
    ("C-23", "Memory Safety Threshold",
     "MEM_SAFETY_THRESHOLD=0.90, MEM_RESERVE_GB=2.0 (configurable)", "done"),
    ("C-25", "Profile Checkpointing",
     "PROFILE_CHECKPOINT_JOBS=10, PROFILE_CHECKPOINT_SECS=30", "done"),
    ("C-27", "CUSTOM/SPLS Scheduler",
     "State-Prioritized Least-Slack: score = slack + memPenalty*1000", "done"),
    ("C-28", "Port Entropy (Round-Robin + TIME_WAIT)",
     "PRE_BASE_PORT=46000 with TIME_WAIT check", "done"),
]

IMPL_MONITOR = [
    ("C-29", "Energy Logging",
     "Energy_uJ column in 27-column CSV (Logger.cpp)", "done"),
    ("C-31", "Centralized Triple Check (FU-4)",
     "popReadyJob() verifies: File readiness, Core gaps (query-based), Memory gaps (SystemMonitor)", "done"),
    ("C-32", "Dual-Phase Welford StdDev (FU-5)",
     "welfordUpdateOn()/welfordUpdatePre() — separate sigma_pre and sigma_on trackers", "done"),
]

IMPL_QA = [
    ("Q-1", "Cap_c Capability Vector Fields",
     "device_class, cpu_freq, ram_mb, net_type, is_default_cap in Job struct", "done"),
    ("Q-2", "TTFR Handshake Probe",
     "Timer 1 (TCP RTT -> theta_est) + Timer 2 (Wrapper -> theta_act)", "done"),
    ("Q-3", "Negotiation State Machine",
     "theta_c >= 1.5 triggers n_alt, 200ms timeout, phi=1 penalty on refusal", "done"),
    ("Q-4", "3-Tier SLO Filtering",
     "Tier 1 (403), Tier 2 (403), Tier 3 VFT (404) with suggested SLO", "done"),
    ("Q-5", "27-Column CSV Trace Schema",
     "+Exec_Time, Speedup_Ratio, Throughput_Local, Original_Model, Completion_Time, Tardiness", "done"),
    ("Q-6", "SA Malleable Core Sizing",
     "4th SA perturbation: adjust_core_count with total_cores bound", "done"),
    ("Q-7", "DQN Hot-Reload",
     "10s file-watch thread with atomic weight swap (DQNScheduler.hpp)", "done"),
    ("Q-8", "Time-Series Memory Telemetry",
     "TELEMETRY_SAMPLE_RATE_MS=100, MEM_TRACE_FILE, telemetryLoop()", "done"),
    ("Q-9", "Strict Core Allocation",
     "acquireCoresStrict() returns exact match or empty vector", "done"),
    ("Q-10", "Asymmetric Burst Time",
     "E_{i,c} = e_pre + theta_c * e_on; theta_c floored at 1.0", "done"),
    ("Q-11", "Psi Utility Metric (Dual-Tier)",
     "Real-time: atomic counters; Post-hoc: analyze_psi.py with Tardiness", "done"),
    ("Q-12", "Poisson Benchmarking Harness",
     "1000 req, lambda={0.1,0.5,1.0,2.0}, 70/20/10 tier mix, D=E_{i,c}*U(1.5,4.0)", "done"),
    ("Q-13", "Scheduling Tournament Script",
     "run_kairos_tournament.sh: all 7 modes, 4 lambda rates, Psi comparison", "done"),
]

PERF_METRICS_EN = [
    ("Completion Time (CT)", "Column 26: Completion_Time = Finish_TS (timestamp when J_on exits)"),
    ("Turnaround Time (TAT)", "Column 11: Actual_TAT = Finish_TS - Arrival_TS"),
    ("Waiting Time (w)", "Column 5: Total_Wait_Time = Start_TS - Arrival_TS"),
    ("Tardiness", "Column 27: Tardiness = max(0, Actual_TAT - Requested_SLO)"),
    ("Success Flag (chi)", "Column 13: SLO_Met = 1 if SLO_Diff >= 0, else 0"),
    ("Exec Time", "Column 22: Exec_Time_ms = Finish_TS - Start_TS (pure processing)"),
    ("Speedup Ratio", "Column 23: Speedup_Ratio = 1/theta_act (>1 = faster than baseline)"),
    ("Throughput (local)", "Column 24: Throughput_Local = 1000/Actual_TAT (per-job)"),
    ("Energy", "Column 9: Energy_uJ from SystemMonitor"),
]

PERF_METRICS_JA = [
    ("完了時刻 (CT)", "カラム 26: Completion_Time = Finish_TS（J_on終了時のタイムスタンプ）"),
    ("ターンアラウンドタイム (TAT)", "カラム 11: Actual_TAT = Finish_TS - Arrival_TS"),
    ("待ち時間 (w)", "カラム 5: Total_Wait_Time = Start_TS - Arrival_TS"),
    ("ターディネス", "カラム 27: Tardiness = max(0, Actual_TAT - Requested_SLO)"),
    ("成功フラグ (chi)", "カラム 13: SLO_Met = 1 (SLO_Diff >= 0の場合), それ以外 0"),
    ("実行時間", "カラム 22: Exec_Time_ms = Finish_TS - Start_TS（純粋な処理時間）"),
    ("スピードアップ比", "カラム 23: Speedup_Ratio = 1/theta_act（>1 = ベースラインより高速）"),
    ("スループット（ローカル）", "カラム 24: Throughput_Local = 1000/Actual_TAT（ジョブ単位）"),
    ("エネルギー", "カラム 9: SystemMonitorからのEnergy_uJ"),
]

TEST_ROWS = [
    (1, "FCFS FIFO Order", "4/4"),
    (2, "SJF Shortest Job First", "4/4"),
    (3, "EDF Earliest Deadline First", "4/4"),
    (4, "LST Least Slack Time", "4/4"),
    (5, "SA Basic Operation", "4/4"),
    (6, "SA SLO-Aware Reordering", "2/2"),
    (7, "SA Memory Pressure", "2/2"),
    (8, "DQN LST Fallback (no weights)", "4/4"),
    (9, "DQN DQN2 Header (64-32-1)", "3/3"),
    (10, "DQN Raw Weights", "2/2"),
    (11, "DQN Legacy/Bad File", "5/5"),
    (12, "Config Parsing (all params incl. Q-3/Q-8/FU-9)", "21/21"),
    (13, "SchedulerFactory All Modes", "8/8"),
    (14, "CUSTOM/SPLS Scheduler", "5/5"),
    (15, "FU-5: Dual-Phase Welford (pre + on)", "12/12"),
    (16, "Q-1: Job Cap_c Fields", "16/16"),
    (17, "Q-6: SA Malleable Core Sizing", "2/2"),
    (18, "Q-7: DQN Hot-Reload", "4/4"),
    (19, "Q-2: TTFR Theta Fields", "6/6"),
    (20, "Q-4: Tier 2 Static SLO Filter", "8/8"),
    (21, "Q-3: Negotiation and Penalty Fields", "9/9"),
    (22, "Config v2: TTFR, Tier 2, Family-Based n_alt, Per-Network TTFR", "8/8"),
    (23, "C-22: Dynamic Watchdog with theta_c", "4/4"),
    (24, "Q-11: Psi Metric Computation", "3/3"),
    (25, "C-17: SA Job Completion Re-Trigger", "2/2"),
    (26, "Contract §4: Completion_Time and Tardiness", "7/7"),
    (27, "Q-12: Poisson SLO Formula D = E_{i,c} * U(1.5, 4.0)", "7/7"),
    (28, "Q-12: 70/20/10 Model Mix Tier Weights", "6/6"),
    (29, "Q-5: Performance Metrics (Throughput, Speedup, Wait, Exec)", "13/13"),
    (30, "FU-3: Amdahl's Law Speedup Model", "8/8"),
    (31, "FU-5: Dual-Phase Welford Independence", "6/6"),
    (32, "FU-9: Per-Network TTFR Baselines", "12/12"),
    (33, "Logger: 27-Column CSV Output Verification", "26/26"),
]

TEST_ROWS_JA = [
    (1, "FCFS — FIFO順序", "4/4"),
    (2, "SJF — 最短ジョブ優先", "4/4"),
    (3, "EDF — 最早デッドライン優先", "4/4"),
    (4, "LST — 最小スラックタイム", "4/4"),
    (5, "SA — 基本操作", "4/4"),
    (6, "SA — SLO考慮並び替え", "2/2"),
    (7, "SA — メモリ圧力", "2/2"),
    (8, "DQN — LSTフォールバック（重みなし）", "4/4"),
    (9, "DQN — DQN2ヘッダーフォーマット (64-32-1)", "3/3"),
    (10, "DQN — 生重みフォーマット", "2/2"),
    (11, "DQN — レガシーDQN1 + 不正ファイル", "5/5"),
    (12, "設定解析 — 全パラメータ (Q-3/Q-8/FU-9含む)", "21/21"),
    (13, "SchedulerFactory — 全7モード", "8/8"),
    (14, "CUSTOM/SPLS — State-Prioritized Least-Slack", "5/5"),
    (15, "FU-5: デュアルフェーズWelford (pre + on)", "12/12"),
    (16, "Q-1: Job Cap_cフィールドとデフォルト", "16/16"),
    (17, "Q-6: SAマリアブルコアサイジング", "2/2"),
    (18, "Q-7: DQNホットリロード", "4/4"),
    (19, "Q-2: TTFR Thetaフィールド", "6/6"),
    (20, "Q-4: Tier 2静的SLOフィルタ", "8/8"),
    (21, "Q-3: ネゴシエーションとペナルティフィールド", "9/9"),
    (22, "設定v2: TTFR, Tier 2, ファミリーベースn_alt, ネットワーク別TTFR", "8/8"),
    (23, "C-22: theta_c対応動的ウォッチドグ", "4/4"),
    (24, "Q-11: Psiメトリクス計算", "3/3"),
    (25, "C-17: SAジョブ完了再トリガー", "2/2"),
    (26, "契約書第4章: Completion_TimeとTardiness", "7/7"),
    (27, "Q-12: ポアソンSLO公式 D = E_{i,c} * U(1.5, 4.0)", "7/7"),
    (28, "Q-12: 70/20/10モデルミックスティア重み", "6/6"),
    (29, "Q-5: パフォーマンスメトリクス (スループット, スピードアップ, 待ち時間, 実行時間)", "13/13"),
    (30, "FU-3: アムダールの法則スピードアップモデル", "8/8"),
    (31, "FU-5: デュアルフェーズWelford独立性", "6/6"),
    (32, "FU-9: ネットワーク別TTFRベースライン", "12/12"),
    (33, "Logger: 27カラムCSV出力検証", "26/26"),
]

FILE_CHANGES = [
    ("Types.hpp", "Cap_c fields, theta/phi, FU-5 dual-phase Welford (pre+on), FU-9 per-network TTFR baselines", "~50 lines added"),
    ("KairosServer.cpp", "FU-4 Triple Check dispatch, FU-5 dual EWMA+Welford, FU-6 family negotiation, FU-7 auto-penalty, FU-9 theta recalc", "Major rewrite"),
    ("KairosServer.hpp", "telemetryLoop() declaration", "1 line added"),
    ("IScheduler.hpp", "FU-4: Centralized Triple Check with ResourceManager* optional param", "~30 lines added"),
    ("ResourceManager.hpp", "acquireCoresStrict() + FU-4 hasFreeInferenceCores() declarations", "2 lines added"),
    ("ResourceManager.cpp", "acquireCoresStrict() + FU-4 hasFreeInferenceCores() implementations", "~30 lines added"),
    ("Logger.cpp", "27-column semicolon CSV with theta/phi/CT/Tardiness columns", "Full rewrite"),
    ("SAScheduler.hpp", "FU-1 TOTAL_CORES=24, FU-3 Amdahl's Law T(p)=e_pre+(theta*e_on)/p^0.8", "Full rewrite"),
    ("DQNScheduler.hpp", "FU-1 normCores=24, Hot-reload thread, loadWeightsInto, atomic swap", "~120 lines added"),
    ("CustomScheduler.hpp", "SPLS: State-Prioritized Least-Slack scheduler", "New file"),
    ("SchedulerFactory.hpp", "SA total_cores parameter, CUSTOM mode", "~5 lines changed"),
    ("ConfigManager.cpp", "FU-6 family-based N_ALT_MAP, FU-9 per-network TTFR parsing", "~30 lines added"),
    ("config.cfg / toy_config.cfg", "FU-1 TOTAL_CORES=24/8, FU-6 family N_ALT_MAP, FU-8 DQN reward weights, FU-9 TTFR baselines", "~25 lines added each"),
    ("scripts/poisson_test_harness.py", "Q-12: 1000 req, lambda CLI, 70/20/10 tier mix, randomized SLO", "Full rewrite"),
    ("scripts/run_kairos_tournament.sh", "Q-13: 7 scheduler modes, 4 lambda rates, Psi comparison", "New file"),
    ("scripts/analyze_psi.py", "Q-11: Psi + Tardiness + Throughput + Speedup + per-model breakdown", "Full rewrite"),
    ("scripts/train_dqn.py", "FU-2 imitation learning bootstrap on EDF, FU-8 multi-objective reward R=w1*SLO+w2*TAT+w3*Energy", "~150 lines"),
    ("src/test_schedulers.cpp", "v6.1: 33 tests, 233 checks — FU-1~FU-10 + Q-1~Q-13 + Contract §4 + Logger CSV", "Full rewrite"),
    ("Howto.md", "Updated binary names (Kairos_server/Kairos_client), Slurm specs, markdown formatting", "Full rewrite"),
    ("toy_run.sh", "Updated client binary name (sappis_client → Kairos_client)", "10 lines changed"),
    (".gitignore", "Updated binary names (Kairos_server, Kairos_client, test_schedulers)", "Rewritten"),
]


# ============================================================
# JA data for implementation tables
# ============================================================

IMPL_CORE_ARCH_JA = [
    ("C-01", "ReLU監査フック（外部ラッパー）",
     "SHARKはブラックボックス; theta_actは外部実行ラッパー（Timer 2）で測定", "done"),
    ("C-02", "Cap_cケーパビリティベクトル",
     "Job構造体: device_class, cpu_freq, ram_mb, net_type, is_default_cap", "done"),
    ("C-03", "3層SLOフィルタリング",
     "Tier 1: D < e_on (403), Tier 2: D < factor*E_{i,c} (403), Tier 3: VFT (404)", "done"),
    ("C-04", "n_altモデルネゴシエーションマッピング (FU-6)",
     "ファミリーベースN_ALT_MAP: alexnet:simc2,hinet; バッチサイズは元リクエストから継承", "done"),
    ("C-05", "差別化EWMA",
     "alpha_pre=0.2（サーバーローカル）, alpha_on=0.4（クライアント依存）", "done"),
    ("C-06", "VFTアドミッションコントロール",
     "VFT_SAFETY_MARGIN=1.2での動的実現可能性チェック", "done"),
    ("C-07", "プロアクティブ前処理バッファ",
     "FileManager: CREATING/READY/IN_USE/USEDライフサイクル", "done"),
]

IMPL_SCHED_ALGO_JA = [
    ("C-08", "DQNバックグラウンド学習パイプライン",
     "train_dqn.py: 5->64->32->1 MLP、DQN2バイナリをエクスポート", "done"),
    ("C-09", "7種のスケジューラモード",
     "FCFS, SJF-Aging, EDF, LST, SA, DQN, CUSTOM/SPLS — SchedulerFactory経由", "done"),
    ("C-10", "Tier 2静的SLOフィルタ",
     "SLO >= SLO_FACTOR_TIER2 * E_{i,c}, E_{i,c} = e_pre + theta_c * e_on", "done"),
    ("C-11", "設定駆動DQN正規化 (FU-1)",
     "TOTAL_CORES=24（エッジノードターゲット）、TOTAL_MEMORY_MBで特徴量正規化", "done"),
    ("C-12", "SA 4D擺動",
     "swap_jobs, adjust_malleable_cores, shift_dwell_time, reverse_subsequence", "done"),
    ("C-13~C-16", "SA多目的コスト関数",
     "Tardiness^2 + 10*MemPenalty + SlackBonus + CoreOvercommit + NegotiationPenalty(phi)", "done"),
    ("C-17", "SAイベントトリガー再最適化",
     "notifyJobComplete() + SA_MIN_RETRIGGER_MS=50（設定可能）", "done"),
    ("C-18", "SAフルキュー可視性",
     "SA_WINDOW_SIZE=100（深さ100までのキュー全体を最適化）", "done"),
    ("C-19", "DQN 5->64->32->1アーキテクチャ",
     "ReLU活性化関数付きPointwise Ranking MLP", "done"),
    ("C-20", "DQN2バイナリフォーマット + CRC32",
     "28バイトヘッダー (magic+dims+arch_hash+crc32) + 生フロート", "done"),
]

IMPL_RESOURCE_JA = [
    ("C-21", "厳密コア割り当て",
     "acquireCoresStrict(): 完全一致または拒否 (ResourceManager.cpp)", "done"),
    ("C-22", "theta_c対応動的ウォッチドグ",
     "timeout_s = min(theta*e_on*1.2/1000+1, MAX_LEASE_TIMEOUT_S)", "done"),
    ("C-23", "メモリ安全閾値",
     "MEM_SAFETY_THRESHOLD=0.90, MEM_RESERVE_GB=2.0（設定可能）", "done"),
    ("C-25", "プロファイルチェックポイント",
     "PROFILE_CHECKPOINT_JOBS=10, PROFILE_CHECKPOINT_SECS=30", "done"),
    ("C-27", "CUSTOM/SPLSスケジューラ",
     "State-Prioritized Least-Slack: score = slack + memPenalty*1000", "done"),
    ("C-28", "ポートエントロピー（ラウンドロビン + TIME_WAIT）",
     "PRE_BASE_PORT=46000 + TIME_WAITチェック", "done"),
]

IMPL_MONITOR_JA = [
    ("C-29", "エネルギーロギング",
     "27カラムCSVのEnergy_uJカラム (Logger.cpp)", "done"),
    ("C-31", "一元化Triple Check (FU-4)",
     "popReadyJob()で検証: ファイル準備、コアギャップ（クエリベース）、メモリギャップ（SystemMonitor）", "done"),
    ("C-32", "デュアルフェーズWelford標準偏差 (FU-5)",
     "welfordUpdateOn()/welfordUpdatePre() — 独立したsigma_preとsigma_onトラッカー", "done"),
]

IMPL_QA_JA = [
    ("Q-1", "Cap_cケーパビリティベクトルフィールド",
     "device_class, cpu_freq, ram_mb, net_type, is_default_cap — Job構造体", "done"),
    ("Q-2", "TTFRハンドシェイクプローブ",
     "Timer 1 (TCP RTT -> theta_est) + Timer 2 (実行ラッパー -> theta_act)", "done"),
    ("Q-3", "ネゴシエーションステートマシン",
     "theta_c >= 1.5でn_alt提案、200msタイムアウト、拒否時phi=1ペナルティ", "done"),
    ("Q-4", "3層SLOフィルタリング",
     "Tier 1 (403), Tier 2 (403), Tier 3 VFT (404) + 推奨SLO", "done"),
    ("Q-5", "27カラムCSVトレーススキーマ",
     "21から27カラムに拡張: +Exec_Time, Speedup_Ratio, Throughput_Local, Original_Model, CT, Tardiness", "done"),
    ("Q-6", "SAマリアブルコアサイジング",
     "SAの第4擺動次元: adjust_core_count + total_cores境界", "done"),
    ("Q-7", "DQNホットリロード",
     "10秒ファイル監視スレッド + アトミック重みスワップ", "done"),
    ("Q-8", "時系列メモリテレメトリ",
     "TELEMETRY_SAMPLE_RATE_MS=100, MEM_TRACE_FILE, telemetryLoop()", "done"),
    ("Q-9", "厳密コア割り当て",
     "acquireCoresStrict(): 完全一致または空ベクトルを返却", "done"),
    ("Q-10", "非対称バーストタイム",
     "E_{i,c} = e_pre + theta_c * e_on; theta_cは1.0でフロア", "done"),
    ("Q-11", "Psiユーティリティメトリクス（デュアルティア）",
     "リアルタイム: アトミックカウンター; ポストホック: analyze_psi.py + Tardiness", "done"),
    ("Q-12", "ポアソンベンチマークハーネス",
     "1000リクエスト, lambda={0.1,0.5,1.0,2.0}, 70/20/10ティアミックス, D=E_{i,c}*U(1.5,4.0)", "done"),
    ("Q-13", "スケジューリングトーナメントスクリプト",
     "run_kairos_tournament.sh: 全7モード、全4ラムダ率、Psi比較レポート", "done"),
]

FILE_CHANGES_JA = [
    ("Types.hpp", "Cap_cフィールド, theta/phi, FU-5デュアルフェーズWelford (pre+on), FU-9ネットワーク別TTFRベースライン", "~50行追加"),
    ("KairosServer.cpp", "FU-4 Triple Checkディスパッチ, FU-5デュアルEWMA+Welford, FU-6ファミリーネゴシエーション, FU-7自動ペナルティ, FU-9 theta再計算", "大規模改修"),
    ("KairosServer.hpp", "telemetryLoop()宣言", "1行追加"),
    ("IScheduler.hpp", "FU-4: ResourceManager*オプションパラメータ付き一元化Triple Check", "~30行追加"),
    ("ResourceManager.hpp", "acquireCoresStrict() + FU-4 hasFreeInferenceCores()宣言", "2行追加"),
    ("ResourceManager.cpp", "acquireCoresStrict() + FU-4 hasFreeInferenceCores()実装", "~30行追加"),
    ("Logger.cpp", "27カラムセミコロンCSV + theta/phi/CT/Tardinessカラム", "全面改修"),
    ("SAScheduler.hpp", "FU-1 TOTAL_CORES=24, FU-3 アムダールの法則 T(p)=e_pre+(theta*e_on)/p^0.8", "全面改修"),
    ("DQNScheduler.hpp", "FU-1 normCores=24, ホットリロードスレッド, loadWeightsInto, アトミックスワップ", "~120行追加"),
    ("CustomScheduler.hpp", "SPLS: State-Prioritized Least-Slackスケジューラ", "新規ファイル"),
    ("SchedulerFactory.hpp", "SA total_coresパラメータ, CUSTOMモード", "~5行変更"),
    ("ConfigManager.cpp", "FU-6ファミリーベースN_ALT_MAP, FU-9ネットワーク別TTFR解析", "~30行追加"),
    ("config.cfg / toy_config.cfg", "FU-1 TOTAL_CORES=24/8, FU-6ファミリーN_ALT_MAP, FU-8 DQN報酬重み, FU-9 TTFRベースライン", "各~25行追加"),
    ("scripts/poisson_test_harness.py", "Q-12: 1000リクエスト, lambda CLI, 70/20/10ティアミックス, ランダムSLO", "全面改修"),
    ("scripts/run_kairos_tournament.sh", "Q-13: 7スケジューラモード, 4ラムダ率, Psi比較", "新規ファイル"),
    ("scripts/analyze_psi.py", "Q-11: Psi + Tardiness + スループット + スピードアップ + モデル別分析", "全面改修"),
    ("scripts/train_dqn.py", "FU-2 EDF模倣学習ブートストラップ, FU-8多目的報酬 R=w1*SLO+w2*TAT+w3*Energy", "~150行"),
    ("src/test_schedulers.cpp", "v6.1: 33テスト, 233チェック — FU-1~FU-10 + Q-1~Q-13 + 契約書§4 + Logger CSV", "全面改修"),
    ("Howto.md", "バイナリ名更新 (Kairos_server/Kairos_client), Slurm仕様, マークダウン整形", "全面改修"),
    ("toy_run.sh", "クライアントバイナリ名更新 (sappis_client → Kairos_client)", "10行変更"),
    (".gitignore", "バイナリ名更新 (Kairos_server, Kairos_client, test_schedulers)", "書き換え"),
]


# ============================================================
# English Document
# ============================================================
def generate_en():
    doc = Document()

    # Title
    doc.add_heading('KairosSNNI Implementation Report & Follow-Up Questions (v3)', level=0)

    p = doc.add_paragraph()
    p.add_run(f"Date: {date.today().isoformat()}").font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run("Reference: Developer_COntract_Updated.pdf (C-01~C-32, Q-1~Q-13) + ClientAnswers4.pdf (FU-1~FU-10)").font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run("Status: Phase 1 Implementation — 233/233 checks passed (v6.1, 33 tests)").font.size = Pt(10)

    # ============================================================
    # SECTION 1: Implementation Status Summary
    # ============================================================
    add_heading_keep(doc, "1. Implementation Status Summary", level=1)
    p = doc.add_paragraph(
        "The following tables summarize the implementation status of all contract items "
        "(C-01 through C-32 and Q-1 through Q-13) as of the current build, including "
        "all FU-1~FU-10 client answers from ClientAnswers4.pdf. "
        "All 233 test checks pass across 33 test categories."
    )
    keep_paragraph_together(p)

    # 1.1 Core Architecture
    add_heading_keep(doc, "1.1 Core Architecture (C-01 ~ C-07)", level=2)
    make_impl_table(doc, ["Item", "Description", "Status"], IMPL_CORE_ARCH)
    doc.add_paragraph("")

    # 1.2 Scheduling Algorithms
    add_heading_keep(doc, "1.2 Scheduling Algorithms (C-08 ~ C-20)", level=2)
    make_impl_table(doc, ["Item", "Description", "Status"], IMPL_SCHED_ALGO)
    doc.add_paragraph("")

    # 1.3 Resource Management
    add_heading_keep(doc, "1.3 Resource Management (C-21 ~ C-28)", level=2)
    make_impl_table(doc, ["Item", "Description", "Status"], IMPL_RESOURCE)
    doc.add_paragraph("")

    # 1.4 Monitoring & Logging
    add_heading_keep(doc, "1.4 Monitoring & Logging (C-29 ~ C-32)", level=2)
    make_impl_table(doc, ["Item", "Description", "Status"], IMPL_MONITOR)
    doc.add_paragraph("")

    # 1.5 Follow-Up Questions Implementation
    add_heading_keep(doc, "1.5 Follow-Up Questions Implementation (Q-1 ~ Q-13)", level=2)
    make_impl_table(doc, ["Item", "Description", "Status"], IMPL_QA)
    doc.add_paragraph("")

    # 1.6 Performance Measurements
    add_heading_keep(doc, "1.6 Performance Measurements (Contract Section 4)", level=2)
    p = doc.add_paragraph(
        "The Logger produces a 27-column semicolon-delimited CSV (scheduler_log.csv) "
        "capturing the following performance metrics per contract Section 4:"
    )
    keep_paragraph_together(p)
    make_metrics_table(doc, PERF_METRICS_EN)

    p = doc.add_paragraph("")
    p = doc.add_paragraph(
        "The post-hoc analyzer (scripts/analyze_psi.py) computes aggregate metrics: "
        "Psi (success density), Throughput (jobs/sec), Speedup (geometric mean), "
        "Waiting Time (mean/P50/P95/P99), Tardiness (mean/P50/P95/P99), Energy total, "
        "and per-model Psi breakdown. Supports --compare for baseline comparison."
    )
    keep_paragraph_together(p)

    # ============================================================
    # SECTION 2: Test Results
    # ============================================================
    add_heading_keep(doc, "2. Test Results", level=1)
    p = doc.add_paragraph(
        "Test suite version: v6.0 — 233/233 checks passed across 33 test categories."
    )
    keep_paragraph_together(p)
    p = doc.add_paragraph(
        "Build command: make test | Source: src/test_schedulers.cpp"
    )
    keep_paragraph_together(p)

    make_test_table(doc, TEST_ROWS)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Total: 33 tests, 233 checks — ALL PASSED")
    run.bold = True
    run.font.size = Pt(11)

    # ============================================================
    # SECTION 3: Technical Implementation Details
    # ============================================================
    add_heading_keep(doc, "3. Technical Implementation Details", level=1)
    p = doc.add_paragraph(
        "Summary of all file changes made based on contract items C-01~C-32, Q-1~Q-13, "
        "and FU-1~FU-10 client answers."
    )
    keep_paragraph_together(p)

    make_file_table(doc, FILE_CHANGES)
    doc.add_paragraph("")

    # ============================================================
    # SECTION 4: FU-1~FU-10 Resolution Summary
    # ============================================================
    add_heading_keep(doc, "4. FU-1~FU-10 Resolution Summary (ClientAnswers4.pdf)", level=1)
    p = doc.add_paragraph(
        "All 10 follow-up questions from the previous report (v2) have been answered by "
        "the client in ClientAnswers4.pdf and fully implemented. Below is a summary of "
        "each resolution and the corresponding code changes."
    )
    keep_paragraph_together(p)

    fu_resolved = [
        ("FU-1", "Target Core Count = 24",
         "TOTAL_CORES=24 in config.cfg; SA bounds clamped to [1, min(16, totalCores)]; DQN normalization updated", "done"),
        ("FU-2", "DQN Imitation Learning Bootstrap",
         "train_dqn.py: Bootstrap on EDF traces (500-1000 episodes) before main training", "done"),
        ("FU-3", "Amdahl's Law Speedup Model",
         "SAScheduler.hpp: T(p) = e_pre + (theta_c * e_on) / p^0.8 replaces linear speedup", "done"),
        ("FU-4", "Centralized Triple Check (C-31)",
         "IScheduler.hpp popReadyJob(): File + Core (query-based) + Memory (SystemMonitor) checks", "done"),
        ("FU-5", "Dual-Phase Welford (C-32)",
         "Types.hpp: Separate welfordUpdateOn()/welfordUpdatePre() with independent counters", "done"),
        ("FU-6", "Family-Based N_ALT_MAP (C-04)",
         "Config format: alexnet:simc2,hinet — batch size preserved from original request", "done"),
        ("FU-7", "Lightweight Model Auto-Penalty",
         "KairosServer.cpp: Models without n_alt entries get phi=1 automatically, skip negotiation", "done"),
        ("FU-8", "Multi-Objective DQN Reward",
         "train_dqn.py: R = (w1*R_SLO) + (w2*R_TAT) + (w3*R_Energy), weights configurable in config.cfg", "done"),
        ("FU-9", "Per-Network TTFR Baselines",
         "LAN=5ms, WiFi=20ms, 5G=40ms, WAN=100ms; theta_est recalculated per net_type", "done"),
        ("FU-10", "Phase 1 Feature Completeness",
         "All contract items C-01~C-32 and Q-1~Q-13 implemented; 33 tests, 233/233 checks passed (incl. Logger CSV verification)", "done"),
    ]
    make_impl_table(doc, ["Item", "Resolution", "Implementation"], fu_resolved)
    doc.add_paragraph("")

    # ============================================================
    # SECTION 5: Remaining Questions for Phase 2 Planning
    # ============================================================
    add_heading_keep(doc, "5. Remaining Questions for Phase 2 Planning", level=1)
    p = doc.add_paragraph(
        "With all FU-1~FU-10 items resolved, the following questions relate to Phase 2 "
        "planning and operational deployment. These are optional items that do not block "
        "Phase 1 delivery but would benefit from client input for future work."
    )
    keep_paragraph_together(p)

    add_heading_keep(doc, "5.1 Deployment & Benchmarking", level=2)

    add_question(doc, 11,
        "Scheduling Tournament Execution Environment",
        "The Poisson test harness (Q-12/Q-13) and tournament script are ready to run. "
        "Should the initial tournament be executed on the 24-core target hardware, or "
        "is a preliminary run on a development node acceptable for Phase 1 delivery? "
        "What is the expected Poisson lambda range for the production workload?",
        context="scripts/run_kairos_tournament.sh, scripts/poisson_test_harness.py",
        priority="Medium"
    )

    add_question(doc, 12,
        "DQN Training Data Volume",
        "The DQN imitation learning bootstrap (FU-2) pre-trains on EDF traces. "
        "For production DQN training, how many scheduling log entries (jobs) should "
        "be accumulated before the first full DQN training cycle? The current pipeline "
        "uses all available data with 80/20 train/test split.",
        context="scripts/train_dqn.py, FU-2",
        priority="Low"
    )

    add_heading_keep(doc, "5.2 Architecture & Extensibility", level=2)

    add_question(doc, 13,
        "Phase Roadmap Beyond Phase 1",
        "The contract references 'Phase 5' (C-01, page 17) but does not define "
        "Phases 2-4. Understanding the roadmap would help ensure Phase 1 "
        "implementation does not create architectural debt for later phases. "
        "Is a formal phase breakdown available?",
        context="C-01, Q-13",
        priority="Low"
    )

    add_question(doc, 14,
        "Amdahl's Law Exponent Calibration",
        "The current Amdahl's Law exponent is k=0.8 (FU-3). This was chosen as a "
        "reasonable default for cryptographic workloads. If empirical per-model "
        "speedup data becomes available from production runs, should the exponent "
        "be calibrated per-model rather than using a global constant?",
        context="SAScheduler.hpp, FU-3",
        priority="Low"
    )

    path = "KairosSNNI_FollowUp_v3_EN.docx"
    doc.save(path)
    print(f"Saved: {path}")
    return path


# ============================================================
# Japanese Document
# ============================================================
def generate_ja():
    doc = Document()

    # Title
    doc.add_heading('KairosSNNI 実装レポート & フォローアップ質問 (v3)', level=0)

    p = doc.add_paragraph()
    p.add_run(f"日付: {date.today().isoformat()}").font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run("参照: Developer_COntract_Updated.pdf (C-01~C-32, Q-1~Q-13) + ClientAnswers4.pdf (FU-1~FU-10)").font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run("ステータス: Phase 1 実装 — 233/233 チェック合格 (v6.1, 33テスト)").font.size = Pt(10)

    # ============================================================
    # SECTION 1: 実装状況
    # ============================================================
    add_heading_keep(doc, "1. 実装状況サマリー", level=1)
    p = doc.add_paragraph(
        "以下の表は、契約の全項目 (C-01～C-32およびQ-1～Q-13) の"
        "現時点での実装状況をまとめたものです。"
        "ClientAnswers4.pdfのFU-1～FU-10も全て実装済みです。"
        "全233テストチェックが33テストカテゴリで合格しています。"
    )
    keep_paragraph_together(p)

    # 1.1 コアアーキテクチャ
    add_heading_keep(doc, "1.1 コアアーキテクチャ (C-01 ~ C-07)", level=2)
    make_impl_table(doc, ["項目", "説明", "ステータス"], IMPL_CORE_ARCH_JA)
    doc.add_paragraph("")

    # 1.2 スケジューリングアルゴリズム
    add_heading_keep(doc, "1.2 スケジューリングアルゴリズム (C-08 ~ C-20)", level=2)
    make_impl_table(doc, ["項目", "説明", "ステータス"], IMPL_SCHED_ALGO_JA)
    doc.add_paragraph("")

    # 1.3 リソース管理
    add_heading_keep(doc, "1.3 リソース管理 (C-21 ~ C-28)", level=2)
    make_impl_table(doc, ["項目", "説明", "ステータス"], IMPL_RESOURCE_JA)
    doc.add_paragraph("")

    # 1.4 モニタリングとロギング
    add_heading_keep(doc, "1.4 モニタリングとロギング (C-29 ~ C-32)", level=2)
    make_impl_table(doc, ["項目", "説明", "ステータス"], IMPL_MONITOR_JA)
    doc.add_paragraph("")

    # 1.5 フォローアップ質問の実装
    add_heading_keep(doc, "1.5 フォローアップ質問の実装 (Q-1 ~ Q-13)", level=2)
    make_impl_table(doc, ["項目", "説明", "ステータス"], IMPL_QA_JA)
    doc.add_paragraph("")

    # 1.6 パフォーマンス測定
    add_heading_keep(doc, "1.6 パフォーマンス測定（契約書第4章）", level=2)
    p = doc.add_paragraph(
        "Loggerは27カラムのセミコロン区切りCSV (scheduler_log.csv) を生成し、"
        "契約書第4章に従った以下のパフォーマンスメトリクスを記録します:"
    )
    keep_paragraph_together(p)
    make_metrics_table(doc, PERF_METRICS_JA)

    p = doc.add_paragraph("")
    p = doc.add_paragraph(
        "ポストホックアナライザー (scripts/analyze_psi.py) はCSVから集約メトリクスを計算します: "
        "Psi（成功密度）、スループット（ジョブ/秒）、スピードアップ（幾何平均）、"
        "待ち時間 (mean/P50/P95/P99)、ターディネス (mean/P50/P95/P99)、エネルギー合計、"
        "モデル別Psi分析。--compareでベースライン比較が可能です。"
    )
    keep_paragraph_together(p)

    # ============================================================
    # SECTION 2: テスト結果
    # ============================================================
    add_heading_keep(doc, "2. テスト結果", level=1)
    p = doc.add_paragraph(
        "テストスイートバージョン: v6.1 — 233/233 チェック合格（32テストカテゴリ）"
    )
    keep_paragraph_together(p)
    p = doc.add_paragraph(
        "ビルドコマンド: make test | ソース: src/test_schedulers.cpp"
    )
    keep_paragraph_together(p)

    make_test_table(doc, TEST_ROWS_JA)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("合計: 33テスト, 233チェック — ALL PASSED")
    run.bold = True
    run.font.size = Pt(11)

    # ============================================================
    # SECTION 3: 技術的実装詳細
    # ============================================================
    add_heading_keep(doc, "3. 技術的実装詳細", level=1)
    p = doc.add_paragraph(
        "契約項目C-01～C-32、Q-1～Q-13、およびFU-1～FU-10に基づく全ファイル変更のサマリー。"
    )
    keep_paragraph_together(p)

    make_file_table(doc, FILE_CHANGES_JA)
    doc.add_paragraph("")

    # ============================================================
    # SECTION 4: FU-1~FU-10 解決サマリー
    # ============================================================
    add_heading_keep(doc, "4. FU-1~FU-10 解決サマリー (ClientAnswers4.pdf)", level=1)
    p = doc.add_paragraph(
        "前回のレポート (v2) の全10項目のフォローアップ質問は、クライアントが"
        "ClientAnswers4.pdfで回答し、全て実装済みです。以下は各解決策と"
        "対応するコード変更のサマリーです。"
    )
    keep_paragraph_together(p)

    fu_resolved_ja = [
        ("FU-1", "ターゲットコア数 = 24",
         "config.cfgでTOTAL_CORES=24; SA探索空間[1, min(16, totalCores)]; DQN正規化更新", "done"),
        ("FU-2", "DQN模倣学習ブートストラップ",
         "train_dqn.py: EDFトレースでブートストラップ（500-1000エピソード）", "done"),
        ("FU-3", "アムダールの法則スピードアップモデル",
         "SAScheduler.hpp: T(p) = e_pre + (theta_c * e_on) / p^0.8 が線形スピードアップを置換", "done"),
        ("FU-4", "一元化Triple Check (C-31)",
         "IScheduler.hpp popReadyJob(): ファイル + コア（クエリベース）+ メモリ（SystemMonitor）チェック", "done"),
        ("FU-5", "デュアルフェーズWelford (C-32)",
         "Types.hpp: 独立したwelfordUpdateOn()/welfordUpdatePre()カウンター", "done"),
        ("FU-6", "ファミリーベースN_ALT_MAP (C-04)",
         "設定フォーマット: alexnet:simc2,hinet — バッチサイズは元リクエストから継承", "done"),
        ("FU-7", "軽量モデル自動ペナルティ",
         "KairosServer.cpp: n_altエントリなしのモデルはphi=1を自動適用、ネゴシエーションスキップ", "done"),
        ("FU-8", "多目的DQN報酬関数",
         "train_dqn.py: R = (w1*R_SLO) + (w2*R_TAT) + (w3*R_Energy), 重みはconfig.cfgで設定可能", "done"),
        ("FU-9", "ネットワーク別TTFRベースライン",
         "LAN=5ms, WiFi=20ms, 5G=40ms, WAN=100ms; theta_estをnet_type別に再計算", "done"),
        ("FU-10", "Phase 1機能完成度",
         "全契約項目C-01~C-32およびQ-1~Q-13実装済み; 33テスト, 233/233チェック合格", "done"),
    ]
    make_impl_table(doc, ["項目", "解決内容", "実装"], fu_resolved_ja)
    doc.add_paragraph("")

    # ============================================================
    # SECTION 5: Phase 2計画のための残課題
    # ============================================================
    add_heading_keep(doc, "5. Phase 2計画のための残課題", level=1)
    p = doc.add_paragraph(
        "FU-1～FU-10の全項目が解決されたため、以下の質問はPhase 2の計画と"
        "運用デプロイメントに関するものです。Phase 1の納品をブロックするものではありませんが、"
        "今後の作業に向けてクライアントからのフィードバックが有益です。"
    )
    keep_paragraph_together(p)

    add_heading_keep(doc, "5.1 デプロイメントとベンチマーク", level=2)

    add_question(doc, 11,
        "スケジューリングトーナメント実行環境",
        "ポアソンテストハーネス (Q-12/Q-13) とトーナメントスクリプトは実行準備が整っています。"
        "初回のトーナメントは24コアのターゲットハードウェアで実行すべきですか、"
        "それともPhase 1の納品には開発ノードでの予備実行で十分ですか？"
        "本番ワークロードで想定されるポアソンラムダの範囲は？",
        context="scripts/run_kairos_tournament.sh, scripts/poisson_test_harness.py",
        priority="中"
    )

    add_question(doc, 12,
        "DQN学習データ量",
        "DQN模倣学習ブートストラップ (FU-2) はEDFトレースで事前学習します。"
        "本番DQN学習のため、最初の完全なDQN学習サイクルの前に"
        "どの程度のスケジューリングログエントリ（ジョブ数）を蓄積すべきですか？"
        "現在のパイプラインは利用可能な全データを80/20の訓練/テスト分割で使用しています。",
        context="scripts/train_dqn.py, FU-2",
        priority="低"
    )

    add_heading_keep(doc, "5.2 アーキテクチャと拡張性", level=2)

    add_question(doc, 13,
        "Phase 1以降のロードマップ",
        "契約書では「Phase 5」(C-01, 17ページ) への言及がありますが、"
        "Phase 2～4は定義されていません。"
        "Phase 1の実装が後のフェーズでアーキテクチャ上の負債を生まないように、"
        "ロードマップの理解が役立ちます。"
        "正式なフェーズ分解はありますか？",
        context="C-01, Q-13",
        priority="低"
    )

    add_question(doc, 14,
        "アムダールの法則指数のキャリブレーション",
        "現在のアムダールの法則の指数はk=0.8 (FU-3) です。"
        "暗号計算ワークロードに対する合理的なデフォルト値として選定されました。"
        "本番運用から実測のモデル別スピードアップデータが得られた場合、"
        "グローバル定数ではなくモデル別に指数をキャリブレーションすべきですか？",
        context="SAScheduler.hpp, FU-3",
        priority="低"
    )

    path = "KairosSNNI_FollowUp_v3_JA.docx"
    doc.save(path)
    print(f"Saved: {path}")
    return path


if __name__ == "__main__":
    generate_en()
    generate_ja()
    print("Done!")
